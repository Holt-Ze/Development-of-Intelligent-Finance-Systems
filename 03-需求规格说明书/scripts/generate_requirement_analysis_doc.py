"""
生成《需求报告分析书》。

输出：
1. 03-需求规格说明书/需求报告分析书.docx
2. 03-需求规格说明书/需求报告分析书.doc

其中 .doc 文件为 HTML 格式的 Word 兼容文档，便于在当前无 Office 自动化组件的环境下稳定生成。
"""

from __future__ import annotations

from datetime import date
from html import escape
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = ROOT / "03-需求规格说明书"
DOCX_OUTPUT = OUTPUT_DIR / "需求报告分析书.docx"
DOC_OUTPUT = OUTPUT_DIR / "需求报告分析书.doc"

TODAY = date(2026, 6, 12)
DATE_TEXT = f"{TODAY.year} 年 {TODAY.month} 月 {TODAY.day} 日"

REFERENCE_DOC = "需求规格.doc（用户提供参考文档）"
PROJECT_NAME = "基于 Web 的学校部门财务管理系统"

ROLE_ROWS = [
    ("登录/退出", "支持", "支持", "统一进入系统并获得角色信息"),
    ("首页仪表盘", "支持", "支持", "查看当月收支汇总、结余和趋势"),
    ("用户管理", "支持", "不支持", "仅管理员维护系统账号与角色"),
    ("部门管理", "支持", "只读/不开放入口", "主数据由管理员统一维护"),
    ("账户管理", "支持", "只读/不开放入口", "保证账户信息口径一致"),
    ("收入类别管理", "支持", "只读/不开放入口", "收入分类由管理员统一维护"),
    ("支出类别管理", "支持", "只读/不开放入口", "支出分类由管理员统一维护"),
    ("收入记录管理", "支持", "支持", "两类角色均可录入与查询收入"),
    ("支出记录管理", "支持", "支持", "两类角色均可录入与查询支出"),
    ("报表中心", "支持", "支持", "周、月、年与自定义区间报表"),
    ("备份管理", "支持", "不支持", "仅管理员可创建、下载和恢复备份"),
]

MODULE_ROWS = [
    ("登录与认证", "LoginPage.vue", "/auth/login, /auth/logout", "sys_user/sys_role/sys_user_role", "返回 token、用户名、真实姓名、角色集合"),
    ("用户管理", "UserPage.vue", "/users", "sys_user/sys_user_role", "管理员维护账号、角色与启停状态"),
    ("部门管理", "DepartmentPage.vue", "/departments", "dept_info", "维护部门编号、名称、负责人、联系电话"),
    ("账户管理", "AccountPage.vue", "/accounts", "account_info", "维护账户编号、名称、开户单位和说明"),
    ("收入类别管理", "IncomeCategoryPage.vue", "/income-categories", "income_category", "收入主数据维护"),
    ("支出类别管理", "ExpenseCategoryPage.vue", "/expense-categories", "expense_category", "支出主数据维护"),
    ("收入记录管理", "IncomePage.vue", "/incomes", "income_record", "编号规则为 IN+日期+三位流水号"),
    ("支出记录管理", "ExpensePage.vue", "/expenses", "expense_record", "编号规则为 EX+日期+三位流水号"),
    ("报表中心", "ReportPage.vue, DashboardPage.vue", "/reports/dashboard, /weekly, /monthly, /yearly, /custom", "income_record/expense_record/dept_info/类别表", "统计有效状态收支数据"),
    ("备份管理", "BackupPage.vue", "/backups, /backups/{id}/download, /backups/{id}/restore", "backup_record", "备份文件命名为 finance_backup_yyyyMMddHHmmss.sql"),
]

ENV_ROWS = [
    ("前端框架", "Vue 3 + Vite + Element Plus + Axios + Vue Router", "单页应用，负责页面展示、路由守卫和 API 调用"),
    ("后端框架", "Spring Boot 3 + MyBatis-Plus", "提供 REST 接口、业务逻辑和数据访问"),
    ("数据库", "MySQL 8", "数据库名为 finance_management_system"),
    ("接口协议", "HTTP + JSON", "统一返回 ApiResponse(code, message, data)"),
    ("运行端口", "前端开发服务默认 5173，后端默认 8080", "前端通过 Axios 调用后端接口"),
    ("浏览器环境", "Chrome / Edge 等现代浏览器", "满足课程设计演示场景"),
]

DATA_ROWS = [
    ("sys_user", "系统用户信息", "username, password, real_name, phone, status", "登录认证、用户状态控制"),
    ("sys_role", "角色定义", "role_code, role_name", "区分 ADMIN 与 FINANCE"),
    ("sys_user_role", "用户角色关系", "user_id, role_id", "实现用户与角色映射"),
    ("dept_info", "部门基础信息", "dept_code, dept_name, leader_name, phone, status", "供收支记录引用"),
    ("account_info", "账户基础信息", "account_code, account_name, owner_unit, description, status", "供收支记录引用"),
    ("income_category", "收入类别主数据", "category_code, category_name, status", "收入业务分类"),
    ("expense_category", "支出类别主数据", "category_code, category_name, status", "支出业务分类"),
    ("income_record", "收入业务记录", "record_no, dept_id, account_id, category_id, amount, occurred_on, operator_name, remark, status", "形成收入明细和报表来源"),
    ("expense_record", "支出业务记录", "record_no, dept_id, account_id, category_id, amount, occurred_on, operator_name, remark, status", "形成支出明细和报表来源"),
    ("backup_record", "备份文件记录", "backup_name, file_path, file_size, created_by, created_at, restored_at, remark", "支撑备份追踪与恢复留痕"),
]

NON_FUNC_ROWS = [
    ("可用性", "界面结构清晰，核心业务流程在少量页面内完成；表单提示明确。", "前端使用 MainLayout、PageCard 和 Element Plus 统一交互样式。"),
    ("性能", "一般列表查询控制在 3 秒内，课程演示数据规模下报表返回平稳。", "报表聚合由 JdbcTemplate 直接执行 SQL 汇总，减少不必要的对象转换。"),
    ("安全性", "未登录用户不得访问业务页面；不同角色仅可访问授权模块。", "前端路由使用 meta.roles 守卫，后端登录采用账号状态校验与密码匹配策略。"),
    ("可维护性", "文档、源码、数据库和建模图分目录归档，命名统一。", "仓库已按课程交付物拆分为 13 个目录，便于验收与答辩。"),
    ("可扩展性", "保留财务导出、预算管理、日志审计等二期扩展空间。", "主数据、业务数据、备份数据独立建表，报表接口已按类型拆分。"),
]

ACCEPTANCE_ROWS = [
    ("角色权限", "管理员可访问用户、主数据、备份页面；财务员仅访问业务与报表页面", "前端路由元信息与登录返回角色集合一致"),
    ("主数据管理", "部门、账户、类别支持新增、修改、启停，不做随意物理删除", "状态字段统一管理引用边界"),
    ("收支管理", "收入与支出记录支持新增、修改、查询、状态维护", "前后端页面、接口、数据表均已对称设计"),
    ("报表统计", "支持周报、月报、年报、自定义区间和首页仪表盘", "对应 /reports 5 组接口"),
    ("数据约束", "金额大于 0，部门/账户/类别需为有效状态", "RecordService 在保存前统一校验"),
    ("备份恢复", "管理员可创建、下载、恢复数据库备份并留痕", "BackupService 写入 backup_record 并记录 restored_at"),
]

KEY_RULES = [
    "系统采用管理员（ADMIN）和财务员（FINANCE）两类角色，登录后按角色返回可访问菜单。",
    "收入记录编号格式为 IN + 发生日期（yyyyMMdd） + 三位流水号；支出记录编号格式为 EX + 发生日期（yyyyMMdd） + 三位流水号。",
    "收入、支出记录在保存前必须校验部门、账户、类别均为启用状态，避免业务数据引用失效主数据。",
    "报表统计仅汇总 status = 1 的有效收支记录，并支持按部门筛选。",
    "系统备份文件采用 SQL 文本形式保存，默认命名为 finance_backup_yyyyMMddHHmmss.sql，并通过 backup_record 表记录元数据。",
    "数据库默认名称为 finance_management_system，后端默认端口为 8080，适合课程设计本地演示环境。",
]

PAIN_POINTS = [
    "纸质台账或分散电子表格难以保证版本一致性，容易出现统计口径不统一的问题。",
    "部门、账户、类别等基础信息往往分散维护，导致收支录入时重复选择、命名不规范。",
    "传统人工汇总周报、月报、年报效率低，且很难做到按部门、按时间区间快速分析。",
    "缺少统一备份与恢复机制时，误删误改后难以回滚，不利于课程演示中的稳定操作。",
]

FUTURE_ROWS = [
    ("导出能力", "增加 Word/PDF 财务分析报告导出，满足归档和上报需求"),
    ("预算管理", "增加预算编制、预算执行对比和预警机制"),
    ("日志审计", "记录关键操作轨迹，提升可追踪性"),
    ("图表增强", "增加更丰富的图表展示与导出能力"),
]


def set_document_defaults(document: DocumentObject) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for style_name, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(3)


def add_paragraph(document: DocumentObject, text: str, *, center: bool = False, bold: bool = False, indent: bool = True) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.first_line_indent = Pt(24) if indent and not center else Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_bullets(document: DocumentObject, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Paragraph")
        paragraph.paragraph_format.left_indent = Pt(24)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_table(document: DocumentObject, headers: list[str], rows: list[tuple[str, ...] | list[str]], widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell(cell, header, bold=True, center=True)
        if widths:
            cell.width = Cm(widths[index])
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell(cells[index], str(value))
            if widths:
                cells[index].width = Cm(widths[index])
    document.add_paragraph()


def set_cell(cell, text: str, *, bold: bool = False, center: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_page_break(document: DocumentObject) -> None:
    document.add_page_break()


def add_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run._r.append(text)
    run._r.append(fld_char3)


def create_cover(document: DocumentObject) -> None:
    for _ in range(7):
        document.add_paragraph()

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("需求报告分析书")
    run.font.name = "Times New Roman"
    run.font.size = Pt(22)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(PROJECT_NAME)
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    document.add_paragraph()
    version = document.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run("V1.0")
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for _ in range(4):
        document.add_paragraph()

    meta_rows = [
        ("课程名称", "软件工程课程设计"),
        ("项目名称", PROJECT_NAME),
        ("文档名称", "需求报告分析书"),
        ("完成日期", DATE_TEXT),
        ("参考资料", REFERENCE_DOC),
        ("备注", "内容已结合当前仓库中的源码、数据库与已有文档整理"),
    ]
    add_table(document, ["项目", "内容"], meta_rows, widths=[4.5, 11.5])


def add_docx_content(document: DocumentObject) -> None:
    document.add_heading("版本记录", level=1)
    add_table(
        document,
        ["版本", "日期", "修改说明"],
        [("1.0", "2026-06-12", "首次生成，结合当前项目源码、数据库和需求说明整理完成")],
        widths=[3, 4, 9],
    )

    document.add_heading("1. 引言", level=1)
    document.add_heading("1.1 编写目的", level=2)
    add_paragraph(document, "本文档用于对“基于 Web 的学校部门财务管理系统”进行需求层面的分析与归纳，明确系统建设背景、业务问题、角色边界、功能范围、数据要求、非功能要求和验收重点，为课程设计中的后续设计、开发、测试和答辩提供统一依据。")
    document.add_heading("1.2 项目背景", level=2)
    add_paragraph(document, "学校内部部门在经费管理过程中，需要持续维护收入、支出、账户、类别、统计报表和数据备份等信息。若长期依赖纸质台账或零散表格，不仅录入效率低，而且容易出现统计口径不一致、历史记录难追踪、信息共享不及时等问题。因此，本项目建设一个轻量、清晰、便于课程演示的 Web 财务管理系统。")
    document.add_heading("1.3 编写依据", level=2)
    add_bullets(
        document,
        [
            "用户提供的参考文档《需求规格.doc》。",
            "仓库中的《需求规格说明书-财务管理系统.md》。",
            "仓库中的《软件设计说明书-财务管理系统.md》与详细设计生成脚本。",
            "前端 Vue 工程、后端 Spring Boot 工程与数据库脚本的当前实现。",
        ],
    )

    document.add_heading("2. 现状与问题分析", level=1)
    document.add_heading("2.1 业务现状", level=2)
    add_paragraph(document, "当前项目的业务场景聚焦于学校部门经费收支管理，涉及部门、账户、收入类别、支出类别等基础信息的维护，以及收入、支出、统计报表和数据库备份等核心流程。项目定位清晰，适合作为软件工程课程设计的完整案例。")
    document.add_heading("2.2 主要痛点", level=2)
    add_bullets(document, PAIN_POINTS)
    document.add_heading("2.3 建设必要性", level=2)
    add_paragraph(document, "建设统一的 Web 财务管理系统后，可将原本分散的基础信息、收支记录和统计工作集中管理，既能提升录入与查询效率，也能保证课程答辩时展示流程完整、数据来源明确、业务规则统一。")

    document.add_heading("3. 项目需求概述", level=1)
    document.add_heading("3.1 建设目标", level=2)
    add_bullets(
        document,
        [
            "统一管理学校部门财务收支相关主数据和业务数据。",
            "支持管理员和财务员按职责完成录入、维护、统计与备份工作。",
            "形成可直接演示的周报、月报、年报、自定义区间报表与首页仪表盘。",
            "降低误操作风险，保留备份记录和恢复能力。",
            "为课程设计归档提供结构完整、内容一致的文档与系统成果。",
        ],
    )
    document.add_heading("3.2 用户角色与权限", level=2)
    add_table(document, ["功能项", "管理员", "财务员", "说明"], ROLE_ROWS, widths=[4, 2.5, 2.5, 7])
    document.add_heading("3.3 运行环境", level=2)
    add_table(document, ["要素", "当前项目内容", "说明"], ENV_ROWS, widths=[3, 6, 7])
    document.add_heading("3.4 功能边界", level=2)
    add_paragraph(document, "本项目当前阶段主要覆盖登录认证、用户管理、部门管理、账户管理、收入类别管理、支出类别管理、收入记录管理、支出记录管理、财务统计与报表、数据备份与恢复等十类核心功能。预算审批、附件报销、多组织租户、审计日志等复杂业务暂未纳入本期范围。")

    document.add_heading("4. 功能需求分析", level=1)
    document.add_heading("4.1 模块与项目实现映射", level=2)
    add_table(document, ["功能模块", "前端页面", "后端接口", "核心数据表", "项目结合点"], MODULE_ROWS, widths=[3, 4, 5, 4, 5])
    document.add_heading("4.2 关键业务规则", level=2)
    add_bullets(document, KEY_RULES)
    document.add_heading("4.3 核心业务流程分析", level=2)
    add_paragraph(document, "系统的主业务流程可概括为“主数据维护 -> 收支录入 -> 报表统计 -> 数据备份”。首先由管理员维护部门、账户和类别等主数据；随后管理员或财务员录入收入、支出信息；系统在保存前校验主数据状态和金额合法性；报表模块再从有效状态记录中汇总出统计结果；最后管理员可在备份管理模块中执行数据库备份、下载或恢复。")

    document.add_heading("5. 数据需求分析", level=1)
    document.add_heading("5.1 核心数据对象", level=2)
    add_table(document, ["表名", "业务含义", "关键字段", "作用说明"], DATA_ROWS, widths=[3.2, 3.5, 6.3, 4.0])
    document.add_heading("5.2 数据关系分析", level=2)
    add_paragraph(document, "从当前数据库脚本可知，收入记录和支出记录均通过外键式业务关联引用部门、账户和类别主数据；用户与角色之间通过 sys_user_role 建立多对多映射；backup_record 单独保存备份文件元数据，以便页面查询与恢复留痕。这样的拆分方式使主数据、业务数据和备份数据边界清楚，便于后续维护和扩展。")
    document.add_heading("5.3 数据约束", level=2)
    add_bullets(
        document,
        [
            "主数据表统一提供 status 字段，用于启用/停用控制。",
            "收入和支出记录必须关联有效的部门、账户和类别。",
            "报表统计仅计算 status = 1 的有效收支记录。",
            "备份恢复需要依赖备份文件真实存在，并记录 restored_at 时间。",
        ],
    )

    document.add_heading("6. 非功能需求分析", level=1)
    add_table(document, ["维度", "需求说明", "当前项目体现"], NON_FUNC_ROWS, widths=[3, 7, 7])

    document.add_heading("7. 技术与接口需求分析", level=1)
    document.add_heading("7.1 前后端协作方式", level=2)
    add_paragraph(document, "前端项目通过 finance.js 对后端接口进行统一封装，接口分组包括 /auth、/users、/departments、/accounts、/income-categories、/expense-categories、/incomes、/expenses、/reports 和 /backups。HTTP 通信层使用 Axios 实例统一附带 token，并对 ApiResponse 结构进行响应解包。")
    document.add_heading("7.2 安全与权限控制", level=2)
    add_paragraph(document, "登录后系统返回 token、用户名、真实姓名和角色集合。前端路由在 beforeEach 守卫中检查 token 是否存在，并依据 meta.roles 判断当前用户是否有权访问目标页面。后端登录服务会校验用户状态，并兼容明文、demo 口径和 BCrypt 密码匹配方式。")
    document.add_heading("7.3 报表与备份的项目特性", level=2)
    add_paragraph(document, "报表模块已实现 dashboard、weekly、monthly、yearly 和 custom 五类统计接口，其中趋势图在较短时间范围内按日聚合、在整年或超过 62 天场景下按月聚合。备份模块则会把 10 张核心表的数据导出为 SQL，并支持下载和恢复，这些都体现了当前项目的完整性和演示价值。")

    document.add_heading("8. 验收建议", level=1)
    add_table(document, ["验收点", "要求", "项目依据"], ACCEPTANCE_ROWS, widths=[3.2, 6.8, 6])

    document.add_heading("9. 风险与后续扩展建议", level=1)
    add_paragraph(document, "从课程设计角度看，当前系统已经具备完整主流程，但若要进一步提升可用性和工程完整度，仍可在导出能力、预算管理、日志审计和图表展示方面继续扩展。建议方向如下：")
    add_table(document, ["方向", "建议内容"], FUTURE_ROWS, widths=[4, 12])

    document.add_heading("10. 结论", level=1)
    add_paragraph(document, "本《需求报告分析书》并非脱离实现单独编写，而是基于当前仓库中的已有需求说明、前后端源码、数据库脚本和文档体系综合整理而成。文档已经将项目中的角色设计、路由页面、接口分组、数据表结构、收支编号规则、报表逻辑和备份机制映射进需求分析结果，能够较好支撑课程设计提交、答辩讲解和后续文档归档。")


def build_docx() -> None:
    document = Document()
    set_document_defaults(document)
    create_cover(document)
    add_page_break(document)
    add_docx_content(document)
    section = document.sections[-1]
    add_page_number(section)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(str(DOCX_OUTPUT))


def html_table(headers: list[str], rows: list[tuple[str, ...] | list[str]]) -> str:
    thead = "".join(f"<th>{escape(header)}</th>" for header in headers)
    body_rows = []
    for row in rows:
        cells = "".join(f"<td>{escape(str(value))}</td>" for value in row)
        body_rows.append(f"<tr>{cells}</tr>")
    return f"<table><thead><tr>{thead}</tr></thead><tbody>{''.join(body_rows)}</tbody></table>"


def html_list(items: list[str]) -> str:
    return "<ul>" + "".join(f"<li>{escape(item)}</li>" for item in items) + "</ul>"


def build_doc_html() -> None:
    parts: list[str] = []
    parts.append("<h1>需求报告分析书</h1>")
    parts.append(f"<h2>{escape(PROJECT_NAME)}</h2>")
    parts.append("<p class='center'>V1.0</p>")
    parts.append(html_table(["项目", "内容"], [
        ("课程名称", "软件工程课程设计"),
        ("项目名称", PROJECT_NAME),
        ("文档名称", "需求报告分析书"),
        ("完成日期", DATE_TEXT),
        ("参考资料", REFERENCE_DOC),
        ("备注", "内容已结合当前仓库中的源码、数据库与已有文档整理"),
    ]))

    parts.append("<h2>版本记录</h2>")
    parts.append(html_table(["版本", "日期", "修改说明"], [
        ("1.0", "2026-06-12", "首次生成，结合当前项目源码、数据库和需求说明整理完成"),
    ]))

    parts.append("<h2>1. 引言</h2>")
    parts.append("<h3>1.1 编写目的</h3>")
    parts.append("<p>本文档用于对“基于 Web 的学校部门财务管理系统”进行需求层面的分析与归纳，明确系统建设背景、业务问题、角色边界、功能范围、数据要求、非功能要求和验收重点，为课程设计中的后续设计、开发、测试和答辩提供统一依据。</p>")
    parts.append("<h3>1.2 项目背景</h3>")
    parts.append("<p>学校内部部门在经费管理过程中，需要持续维护收入、支出、账户、类别、统计报表和数据备份等信息。若长期依赖纸质台账或零散表格，不仅录入效率低，而且容易出现统计口径不一致、历史记录难追踪、信息共享不及时等问题。因此，本项目建设一个轻量、清晰、便于课程演示的 Web 财务管理系统。</p>")
    parts.append("<h3>1.3 编写依据</h3>")
    parts.append(html_list([
        "用户提供的参考文档《需求规格.doc》。",
        "仓库中的《需求规格说明书-财务管理系统.md》。",
        "仓库中的《软件设计说明书-财务管理系统.md》与详细设计生成脚本。",
        "前端 Vue 工程、后端 Spring Boot 工程与数据库脚本的当前实现。",
    ]))

    parts.append("<h2>2. 现状与问题分析</h2>")
    parts.append("<h3>2.1 业务现状</h3>")
    parts.append("<p>当前项目的业务场景聚焦于学校部门经费收支管理，涉及部门、账户、收入类别、支出类别等基础信息的维护，以及收入、支出、统计报表和数据库备份等核心流程。项目定位清晰，适合作为软件工程课程设计的完整案例。</p>")
    parts.append("<h3>2.2 主要痛点</h3>")
    parts.append(html_list(PAIN_POINTS))
    parts.append("<h3>2.3 建设必要性</h3>")
    parts.append("<p>建设统一的 Web 财务管理系统后，可将原本分散的基础信息、收支记录和统计工作集中管理，既能提升录入与查询效率，也能保证课程答辩时展示流程完整、数据来源明确、业务规则统一。</p>")

    parts.append("<h2>3. 项目需求概述</h2>")
    parts.append("<h3>3.1 建设目标</h3>")
    parts.append(html_list([
        "统一管理学校部门财务收支相关主数据和业务数据。",
        "支持管理员和财务员按职责完成录入、维护、统计与备份工作。",
        "形成可直接演示的周报、月报、年报、自定义区间报表与首页仪表盘。",
        "降低误操作风险，保留备份记录和恢复能力。",
        "为课程设计归档提供结构完整、内容一致的文档与系统成果。",
    ]))
    parts.append("<h3>3.2 用户角色与权限</h3>")
    parts.append(html_table(["功能项", "管理员", "财务员", "说明"], ROLE_ROWS))
    parts.append("<h3>3.3 运行环境</h3>")
    parts.append(html_table(["要素", "当前项目内容", "说明"], ENV_ROWS))
    parts.append("<h3>3.4 功能边界</h3>")
    parts.append("<p>本项目当前阶段主要覆盖登录认证、用户管理、部门管理、账户管理、收入类别管理、支出类别管理、收入记录管理、支出记录管理、财务统计与报表、数据备份与恢复等十类核心功能。预算审批、附件报销、多组织租户、审计日志等复杂业务暂未纳入本期范围。</p>")

    parts.append("<h2>4. 功能需求分析</h2>")
    parts.append("<h3>4.1 模块与项目实现映射</h3>")
    parts.append(html_table(["功能模块", "前端页面", "后端接口", "核心数据表", "项目结合点"], MODULE_ROWS))
    parts.append("<h3>4.2 关键业务规则</h3>")
    parts.append(html_list(KEY_RULES))
    parts.append("<h3>4.3 核心业务流程分析</h3>")
    parts.append("<p>系统的主业务流程可概括为“主数据维护 - 收支录入 - 报表统计 - 数据备份”。首先由管理员维护部门、账户和类别等主数据；随后管理员或财务员录入收入、支出信息；系统在保存前校验主数据状态和金额合法性；报表模块再从有效状态记录中汇总出统计结果；最后管理员可在备份管理模块中执行数据库备份、下载或恢复。</p>")

    parts.append("<h2>5. 数据需求分析</h2>")
    parts.append("<h3>5.1 核心数据对象</h3>")
    parts.append(html_table(["表名", "业务含义", "关键字段", "作用说明"], DATA_ROWS))
    parts.append("<h3>5.2 数据关系分析</h3>")
    parts.append("<p>从当前数据库脚本可知，收入记录和支出记录均通过外键式业务关联引用部门、账户和类别主数据；用户与角色之间通过 sys_user_role 建立多对多映射；backup_record 单独保存备份文件元数据，以便页面查询与恢复留痕。这样的拆分方式使主数据、业务数据和备份数据边界清楚，便于后续维护和扩展。</p>")
    parts.append("<h3>5.3 数据约束</h3>")
    parts.append(html_list([
        "主数据表统一提供 status 字段，用于启用/停用控制。",
        "收入和支出记录必须关联有效的部门、账户和类别。",
        "报表统计仅计算 status = 1 的有效收支记录。",
        "备份恢复需要依赖备份文件真实存在，并记录 restored_at 时间。",
    ]))

    parts.append("<h2>6. 非功能需求分析</h2>")
    parts.append(html_table(["维度", "需求说明", "当前项目体现"], NON_FUNC_ROWS))

    parts.append("<h2>7. 技术与接口需求分析</h2>")
    parts.append("<h3>7.1 前后端协作方式</h3>")
    parts.append("<p>前端项目通过 finance.js 对后端接口进行统一封装，接口分组包括 /auth、/users、/departments、/accounts、/income-categories、/expense-categories、/incomes、/expenses、/reports 和 /backups。HTTP 通信层使用 Axios 实例统一附带 token，并对 ApiResponse 结构进行响应解包。</p>")
    parts.append("<h3>7.2 安全与权限控制</h3>")
    parts.append("<p>登录后系统返回 token、用户名、真实姓名和角色集合。前端路由在 beforeEach 守卫中检查 token 是否存在，并依据 meta.roles 判断当前用户是否有权访问目标页面。后端登录服务会校验用户状态，并兼容明文、demo 口径和 BCrypt 密码匹配方式。</p>")
    parts.append("<h3>7.3 报表与备份的项目特性</h3>")
    parts.append("<p>报表模块已实现 dashboard、weekly、monthly、yearly 和 custom 五类统计接口，其中趋势图在较短时间范围内按日聚合、在整年或超过 62 天场景下按月聚合。备份模块则会把 10 张核心表的数据导出为 SQL，并支持下载和恢复，这些都体现了当前项目的完整性和演示价值。</p>")

    parts.append("<h2>8. 验收建议</h2>")
    parts.append(html_table(["验收点", "要求", "项目依据"], ACCEPTANCE_ROWS))

    parts.append("<h2>9. 风险与后续扩展建议</h2>")
    parts.append("<p>从课程设计角度看，当前系统已经具备完整主流程，但若要进一步提升可用性和工程完整度，仍可在导出能力、预算管理、日志审计和图表展示方面继续扩展。建议方向如下：</p>")
    parts.append(html_table(["方向", "建议内容"], FUTURE_ROWS))

    parts.append("<h2>10. 结论</h2>")
    parts.append("<p>本《需求报告分析书》并非脱离实现单独编写，而是基于当前仓库中的已有需求说明、前后端源码、数据库脚本和文档体系综合整理而成。文档已经将项目中的角色设计、路由页面、接口分组、数据表结构、收支编号规则、报表逻辑和备份机制映射进需求分析结果，能够较好支撑课程设计提交、答辩讲解和后续文档归档。</p>")

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8" />
  <title>需求报告分析书</title>
  <style>
    body {{
      font-family: "SimSun", "宋体", serif;
      margin: 32px 40px;
      line-height: 1.75;
      color: #111;
    }}
    h1, h2, h3 {{
      font-family: "SimHei", "黑体", sans-serif;
      color: #111;
    }}
    h1 {{
      text-align: center;
      font-size: 28px;
      margin-bottom: 8px;
    }}
    h2 {{
      font-size: 20px;
      margin-top: 28px;
    }}
    h3 {{
      font-size: 16px;
      margin-top: 18px;
    }}
    p {{
      text-indent: 2em;
      margin: 8px 0;
    }}
    p.center {{
      text-indent: 0;
      text-align: center;
    }}
    ul {{
      margin: 8px 0 8px 24px;
    }}
    li {{
      margin: 4px 0;
    }}
    table {{
      border-collapse: collapse;
      width: 100%;
      margin: 12px 0 18px;
      table-layout: fixed;
      word-break: break-word;
    }}
    th, td {{
      border: 1px solid #333;
      padding: 8px 10px;
      font-size: 14px;
      vertical-align: top;
    }}
    th {{
      background: #f3f3f3;
      text-align: center;
    }}
  </style>
</head>
<body>
  {''.join(parts)}
</body>
</html>
"""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    DOC_OUTPUT.write_text(html, encoding="utf-8")


def main() -> None:
    build_docx()
    build_doc_html()
    print(f"Generated: {DOCX_OUTPUT}")
    print(f"Generated: {DOC_OUTPUT}")


if __name__ == "__main__":
    main()
