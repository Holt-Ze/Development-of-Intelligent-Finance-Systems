from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
PROJECT_NAME = "基于 Web 的学校部门财务管理系统"
DATE_TEXT = "2026 年 6 月 19 日"


@dataclass(frozen=True)
class DocSpec:
    title: str
    markdown_path: Path
    output_path: Path


DOCS = [
    DocSpec(
        title="概要设计",
        markdown_path=ROOT / "04-软件设计说明书" / "概要设计-财务管理系统.md",
        output_path=ROOT / "04-软件设计说明书" / "概要设计-财务管理系统.docx",
    ),
    DocSpec(
        title="测试计划",
        markdown_path=ROOT / "05-测试分析报告" / "测试计划-财务管理系统.md",
        output_path=ROOT / "05-测试分析报告" / "测试计划-财务管理系统.docx",
    ),
    DocSpec(
        title="测试设计",
        markdown_path=ROOT / "05-测试分析报告" / "测试设计-财务管理系统.md",
        output_path=ROOT / "05-测试分析报告" / "测试设计-财务管理系统.docx",
    ),
    DocSpec(
        title="测试跟踪日志",
        markdown_path=ROOT / "05-测试分析报告" / "测试跟踪日志-财务管理系统.md",
        output_path=ROOT / "05-测试分析报告" / "测试跟踪日志-财务管理系统.docx",
    ),
    DocSpec(
        title="提交手册",
        markdown_path=ROOT / "14-提交手册" / "提交手册-财务管理系统.md",
        output_path=ROOT / "14-提交手册" / "提交手册-财务管理系统.docx",
    ),
]


def set_document_defaults(document: DocumentObject) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for style_name, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(3)


def set_run_font(run, size: float = 12, bold: bool = False, east_asia: str = "宋体") -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def add_cover(document: DocumentObject, title: str) -> None:
    lines = ["软件工程课程设计文档", PROJECT_NAME, title, "版本：V1.0", DATE_TEXT]
    for index, text in enumerate(lines):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if index == 0:
            paragraph.paragraph_format.space_before = Pt(120)
        elif index == 1:
            paragraph.paragraph_format.space_before = Pt(36)
        else:
            paragraph.paragraph_format.space_before = Pt(18)
        run = paragraph.add_run(text)
        if index == 0:
            set_run_font(run, size=18, bold=True, east_asia="黑体")
        elif index in (1, 2):
            set_run_font(run, size=20, bold=True, east_asia="黑体")
        else:
            set_run_font(run, size=14, bold=False, east_asia="宋体")


def add_toc(paragraph) -> None:
    fld_simple = OxmlElement("w:fldSimple")
    fld_simple.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "目录将在打开文档后自动更新"
    run.append(text)
    fld_simple.append(run)
    paragraph._p.append(fld_simple)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run._r.append(text)
    run._r.append(fld_char3)


def restart_page_numbering(section) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), "1")


def set_footer_page_number(section) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.text = ""
    add_page_number(paragraph)


def add_body_paragraph(document: DocumentObject, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run)


def add_list_paragraph(document: DocumentObject, text: str, prefix: str = "• ") -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(24)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(f"{prefix}{text}")
    set_run_font(run)


def add_code_block(document: DocumentObject, code_lines: list[str]) -> None:
    for line in code_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.left_indent = Pt(24)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
    document.add_paragraph()


def set_cell_text(cell, text: str, bold: bool = False, center: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(document: DocumentObject, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, cell_text in enumerate(rows[0]):
        set_cell_text(table.rows[0].cells[index], cell_text, bold=True, center=True)
    for row in rows[1:]:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
    document.add_paragraph()


def is_special_line(line: str) -> bool:
    stripped = line.strip()
    return (
        not stripped
        or stripped.startswith("#")
        or stripped.startswith("|")
        or stripped.startswith("```")
        or bool(re.match(r"^[-*]\s+", stripped))
        or bool(re.match(r"^\d+\.\s+", stripped))
    )


def parse_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def is_separator_row(row: list[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", cell) for cell in row)


def render_markdown(document: DocumentObject, markdown_text: str) -> None:
    lines = markdown_text.splitlines()
    index = 0
    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            document.add_heading(heading_match.group(2).strip(), level=level)
            index += 1
            continue

        if stripped.startswith("```"):
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index].rstrip("\n"))
                index += 1
            add_code_block(document, code_lines)
            index += 1
            continue

        if stripped.startswith("|"):
            table_lines: list[list[str]] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                row = parse_table_row(lines[index])
                if row and not is_separator_row(row):
                    table_lines.append(row)
                index += 1
            add_table(document, table_lines)
            continue

        bullet_match = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet_match:
            while index < len(lines):
                item_match = re.match(r"^[-*]\s+(.*)$", lines[index].strip())
                if not item_match:
                    break
                add_list_paragraph(document, item_match.group(1).strip())
                index += 1
            continue

        numbered_match = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered_match:
            while index < len(lines):
                item_match = re.match(r"^(\d+)\.\s+(.*)$", lines[index].strip())
                if not item_match:
                    break
                add_list_paragraph(document, item_match.group(2).strip(), prefix=f"{item_match.group(1)}. ")
                index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines) and not is_special_line(lines[index]):
            paragraph_lines.append(lines[index].strip())
            index += 1
        add_body_paragraph(document, "".join(paragraph_lines))


def build_doc(spec: DocSpec) -> None:
    markdown_text = spec.markdown_path.read_text(encoding="utf-8")
    document = Document()
    set_document_defaults(document)
    add_cover(document, spec.title)

    content_section = document.add_section(WD_SECTION_START.NEW_PAGE)
    restart_page_numbering(content_section)
    set_footer_page_number(content_section)

    toc_title = document.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run("目录")
    set_run_font(run, size=16, bold=True, east_asia="黑体")
    toc_paragraph = document.add_paragraph()
    add_toc(toc_paragraph)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    render_markdown(document, markdown_text)

    spec.output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(spec.output_path)


def main() -> None:
    for spec in DOCS:
        build_doc(spec)
        print(f"generated: {spec.output_path.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
