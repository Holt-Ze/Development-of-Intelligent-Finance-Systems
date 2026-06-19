"""
Generate the Detailed Design Document (详细设计说明书) for the Financial Management System.
Follows the reference document structure (详细设计.doc) adapted for Spring Boot 3 + Vue 3 architecture.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph
from docx.table import Table

ROOT = Path(__file__).resolve().parents[2]
OUTPUT_PATH = ROOT / '04-软件设计说明书' / '详细设计-财务管理系统.docx'


# ── helpers ──────────────────────────────────────────────────────────

def set_document_defaults(document: DocumentObject) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    styles = document.styles
    normal = styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)

    heading1 = styles['Heading 1']
    heading1.font.name = 'Times New Roman'
    heading1.font.bold = True
    heading1.font.size = Pt(16)
    heading1.font.color.rgb = RGBColor(0, 0, 0)
    heading1.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    heading1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    heading1.paragraph_format.first_line_indent = Pt(0)
    heading1.paragraph_format.space_before = Pt(12)
    heading1.paragraph_format.space_after = Pt(6)

    for style_name, size in [('Heading 2', 14), ('Heading 3', 12)]:
        style = styles[style_name]
        style.font.name = 'Times New Roman'
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(3)


def set_cell_text(cell, text: str, bold: bool = False, center: bool = False) -> None:
    cell.text = ''
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)
    run.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders_none(table: Table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.append(tbl_pr)
    tbl_borders = tbl_pr.first_child_found_in('w:tblBorders')
    if tbl_borders is None:
        tbl_borders = OxmlElement('w:tblBorders')
        tbl_pr.append(tbl_borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = tbl_borders.find(qn(f'w:{edge}'))
        if element is None:
            element = OxmlElement(f'w:{edge}')
            tbl_borders.append(element)
        element.set(qn('w:val'), 'nil')


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in('w:tcBorders')
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)
    for edge, edge_data in kwargs.items():
        tag = qn(f'w:{edge}')
        element = tc_borders.find(tag)
        if element is None:
            element = OxmlElement(f'w:{edge}')
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f'w:{key}'), str(value))


def apply_three_line_table(table: Table, header_rows: int = 1) -> None:
    set_table_borders_none(table)
    if not table.rows:
        return
    top_border = {'val': 'single', 'sz': '10', 'color': '000000'}
    mid_border = {'val': 'single', 'sz': '6', 'color': '000000'}
    bottom_border = {'val': 'single', 'sz': '10', 'color': '000000'}
    first_row = table.rows[0]
    for cell in first_row.cells:
        set_cell_border(cell, top=top_border)
    header_row = table.rows[header_rows - 1]
    for cell in header_row.cells:
        set_cell_border(cell, bottom=mid_border)
    last_row = table.rows[-1]
    for cell in last_row.cells:
        set_cell_border(cell, bottom=bottom_border)


def style_body_paragraph(paragraph: Paragraph, align: WD_ALIGN_PARAGRAPH | None = None) -> None:
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    if align is not None:
        paragraph.alignment = align


def add_body_paragraph(document: DocumentObject, text: str,
                       align: WD_ALIGN_PARAGRAPH | None = None) -> Paragraph:
    paragraph = document.add_paragraph()
    paragraph.add_run(text)
    style_body_paragraph(paragraph, align)
    return paragraph


def add_bullet_paragraph(document: DocumentObject, text: str) -> Paragraph:
    paragraph = document.add_paragraph(style='List Paragraph')
    paragraph.add_run(text)
    paragraph.paragraph_format.left_indent = Pt(24)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_after = Pt(0)
    return paragraph


def add_simple_table(document: DocumentObject, headers: list[str],
                     rows: list[tuple[str, ...]], caption: str) -> Table:
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    cap_run.font.name = 'Times New Roman'
    cap_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    cap_run.font.size = Pt(10.5)
    cap_run.bold = True

    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, center=True)
    for row in rows:
        row_cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(row_cells[idx], value, center=False)
    apply_three_line_table(table)
    document.add_paragraph()
    return table


def add_page_break(document: DocumentObject) -> None:
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_toc(paragraph: Paragraph) -> None:
    fld_simple = OxmlElement('w:fldSimple')
    fld_simple.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
    run = OxmlElement('w:r')
    text = OxmlElement('w:t')
    text.text = '目录将在打开文档后自动更新'
    run.append(text)
    fld_simple.append(run)
    paragraph._p.append(fld_simple)


def add_page_number(paragraph: Paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'PAGE'
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'separate')
    text = OxmlElement('w:t')
    text.text = '1'
    fld_char3 = OxmlElement('w:fldChar')
    fld_char3.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run._r.append(text)
    run._r.append(fld_char3)


def restart_page_numbering(section) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn('w:pgNumType'))
    if pg_num_type is None:
        pg_num_type = OxmlElement('w:pgNumType')
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn('w:start'), '1')


def set_footer_page_number(section) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.text = ''
    add_page_number(paragraph)


def add_code_block(document: DocumentObject, code_lines: list[str], caption: str = '') -> None:
    if caption:
        cap = document.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(caption)
        cap_run.font.name = 'Times New Roman'
        cap_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        cap_run.font.size = Pt(10.5)
        cap_run.bold = True
    for line in code_lines:
        p = document.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Pt(24)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    document.add_paragraph()


# ── cover ────────────────────────────────────────────────────────────

def create_cover(document: DocumentObject) -> None:
    for _ in range(8):
        document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('财务管理系统详细设计')
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.bold = True
    run.font.size = Pt(22)

    document.add_paragraph()
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('基于 Web 的学校部门财务管理系统')
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(16)

    document.add_paragraph()
    version = document.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run('V1.0')
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(16)

    for _ in range(4):
        document.add_paragraph()

    meta = [
        ('课程名称', '软件工程课程设计'),
        ('项目名称', '基于 Web 的学校部门财务管理系统'),
        ('完成日期', '2026 年 6 月 12 日'),
        ('姓名', '【待填写】'),
        ('学号', '【待填写】'),
        ('学院', '【待填写】'),
        ('专业', '【待填写】'),
        ('指导教师', '【待填写】'),
    ]
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for key, value in meta:
        cells = table.add_row().cells
        set_cell_text(cells[0], key, bold=True, center=True)
        set_cell_text(cells[1], value)


def add_version_record(document: DocumentObject) -> None:
    document.add_heading('版本记录', level=1)
    add_simple_table(
        document,
        ['修改编号', '修改日期', '修改后版本', '修改位置', '修改内容概述'],
        [('000', '2026.06.12', '1.0', '全部', '初始发布版本')],
        '表 0-1 版本更新记录',
    )


# ── Section 1: Introduction ─────────────────────────────────────────

def add_introduction(document: DocumentObject) -> None:
    document.add_heading('1. 导言', level=1)

    document.add_heading('1.1 目的', level=2)
    add_body_paragraph(
        document,
        '该文档的目的是描述《基于 Web 的学校部门财务管理系统》项目的详细设计，其主要内容包括：'
    )
    add_bullet_paragraph(document, '系统功能简介')
    add_bullet_paragraph(document, '系统详细设计简述')
    add_bullet_paragraph(document, '各个模块的三层划分（表示层、控制层、模型层）')
    add_bullet_paragraph(document, '核心模块组件的伪代码描述')
    add_body_paragraph(document, '本文档的预期读者是：开发人员、项目管理人员、测试人员。')

    document.add_heading('1.2 范围', level=2)
    add_body_paragraph(
        document,
        '该文档定义了系统的各个模块和模块接口，描述了每个模块在视图层（前端 Vue 页面与组件）、'
        '控制层（Spring MVC Controller）和模型层（Service、Mapper、Entity）的详细设计。'
        '未确定单元的具体实现细节将在编码实现阶段进一步细化。'
    )

    document.add_heading('1.3 缩写说明', level=2)
    add_simple_table(
        document,
        ['缩写', '全称', '说明'],
        [
            ('MVC', 'Model-View-Controller', '模式－视图－控制，表示三层结构体系'),
            ('SPA', 'Single Page Application', '单页应用，前端通过路由动态切换视图'),
            ('REST', 'Representational State Transfer', '表述性状态转移，API 接口风格'),
            ('ORM', 'Object-Relational Mapping', '对象关系映射，本系统使用 MyBatis-Plus'),
            ('DTO', 'Data Transfer Object', '数据传输对象，用于接收请求参数'),
            ('VO', 'View Object', '视图对象，用于封装返回给前端的数据'),
            ('BCrypt', 'Blowfish Cipher', '密码加密算法，用于用户密码安全存储'),
            ('SQL', 'Structured Query Language', '结构化查询语言'),
            ('JWT', 'JSON Web Token', '本系统使用 UUID Token 替代标准 JWT'),
        ],
        '表 1-1 缩写说明',
    )

    document.add_heading('1.4 术语定义', level=2)
    add_simple_table(
        document,
        ['术语', '定义'],
        [
            ('Spring Boot', '基于 Spring 框架的快速应用开发框架，提供自动配置与起步依赖'),
            ('MyBatis-Plus', 'MyBatis 的增强工具，提供内置 CRUD 方法与条件构造器'),
            ('Vue 3', '渐进式 JavaScript 前端框架，采用组合式 API'),
            ('Element Plus', '基于 Vue 3 的桌面端 UI 组件库'),
            ('Vite', '新一代前端构建工具，提供极速开发服务与打包能力'),
            ('Axios', '基于 Promise 的 HTTP 客户端，用于前后端数据交互'),
            ('Vue Router', 'Vue 官方路由管理器，控制单页应用的页面导航和权限守卫'),
            ('管理员', '负责系统用户、部门、账户、类别与备份管理的角色（角色编码 ADMIN）'),
            ('财务员', '负责收入、支出录入以及报表查询的业务角色（角色编码 FINANCE）'),
            ('主数据', '部门、账户、收入类别、支出类别等被业务记录引用的基础信息'),
            ('有效状态', '状态值为 1，可参与业务录入、查询与统计的记录'),
            ('停用状态', '状态值为 0，数据保留但不再允许被新增业务引用'),
        ],
        '表 1-2 术语定义',
    )

    document.add_heading('1.5 引用标准', level=2)
    add_bullet_paragraph(document, '[1] 《软件工程课程设计文档规范》 V1.0')
    add_bullet_paragraph(document, '[2] 《软件详细设计报告格式标准》 V1.0')

    document.add_heading('1.6 参考资料', level=2)
    add_simple_table(
        document,
        ['编号', '参考资料'],
        [
            ('[1]', 'README.md（项目总体说明与运行信息）'),
            ('[2]', '需求规格说明书-财务管理系统.md'),
            ('[3]', '概要设计-财务管理系统.docx'),
            ('[4]', '软件设计说明书-财务管理系统.md'),
            ('[5]', '财务管理系统.sql（数据库建表脚本）'),
            ('[6]', 'backend/src/main/resources/application.yml（后端配置）'),
            ('[7]', 'Vue 3 官方文档（https://cn.vuejs.org/）'),
            ('[8]', 'Spring Boot 3 官方文档（https://docs.spring.io/spring-boot/）'),
            ('[9]', 'MyBatis-Plus 官方文档（https://baomidou.com/）'),
        ],
        '表 1-3 参考资料',
    )

    document.add_heading('1.7 版本更新信息', level=2)
    add_body_paragraph(
        document,
        '本文档版本更新记录见表 0-1。V1.0 为初始发布版本，依据当前仓库中的需求规格说明、'
        '概要设计、数据库脚本、运行配置与前后端实现整理而成，面向课程设计归档与答辩场景使用。'
    )


# ── Section 2: System Design Overview ───────────────────────────────

def add_system_overview(document: DocumentObject) -> None:
    document.add_heading('2. 系统设计概述', level=1)

    add_body_paragraph(
        document,
        '根据《学校部门财务管理系统》的概要设计，本系统采用前后端分离的 B/S 架构，从功能角度分解，'
        '可以分为客户端子系统和管理端子系统。管理端系统包含登录与权限管理、用户管理、部门管理、'
        '账户管理、收入类别管理、支出类别管理、收入记录管理、支出记录管理、财务统计与报表、'
        '数据备份与恢复共 10 个功能模块。各模块通过统一的 REST 接口层进行前后端协作。'
    )

    add_body_paragraph(
        document,
        '系统前端基于 Vue 3 + Element Plus 构建单页应用（SPA），通过 Vue Router 管理页面导航与'
        '权限守卫，通过 Axios 封装 HTTP 请求并与后端进行数据交互。系统后端基于 Spring Boot 3 + '
        'MyBatis-Plus 构建 RESTful API 服务，提供统一的 JSON 响应格式和全局异常处理。数据库采用 '
        'MySQL 8，使用 InnoDB 存储引擎与 utf8mb4 字符编码。'
    )

    add_simple_table(
        document,
        ['功能模块', '主要职责', '对应子系统'],
        [
            ('登录与权限管理', '用户认证、Token 下发、角色识别与菜单控制', '管理端'),
            ('用户管理', '系统用户的增删改查与状态管理', '管理端'),
            ('部门管理', '部门基础信息维护', '管理端'),
            ('账户管理', '财务账户基础信息维护', '管理端'),
            ('收入类别管理', '收入类别的维护与状态控制', '管理端'),
            ('支出类别管理', '支出类别的维护与状态控制', '管理端'),
            ('收入记录管理', '收入业务记录的录入、修改、查询与停用', '管理端'),
            ('支出记录管理', '支出业务记录的录入、修改、查询与停用', '管理端'),
            ('财务统计与报表', '周/月/年/自定义统计与仪表盘', '管理端'),
            ('数据备份与恢复', '数据库备份创建、下载与恢复', '管理端'),
            ('客户端系统', '前端 SPA 应用整体框架、路由与组件', '客户端'),
        ],
        '表 2-1 系统模块划分',
    )


# ── Section 3: Detailed Design Overview ─────────────────────────────

def add_detailed_design_overview(document: DocumentObject) -> None:
    document.add_heading('3. 详细设计概述', level=1)

    add_body_paragraph(
        document,
        '由于本系统采用了基于 Spring Boot + Vue 3 的前后端分离架构，后端采用 MVC 三层设计模式'
        '（Controller 控制层、Service 业务层、Mapper 数据访问层），前端采用 MVVM 组件化设计模式'
        '（View 视图组件、Router 路由控制、API 数据层）。在整个开发过程中，尽可能采用复用的原则，'
        '例如统一封装 Axios 请求模块、统一 API 响应格式、统一全局异常处理等。'
    )

    add_body_paragraph(
        document,
        '本文档的详细设计按照前端视图层（Vue 页面与组件）、后端控制层（Controller）、'
        '后端模型层（Service + Mapper + Entity）三个层次分别进行描述，并对核心模块给出伪代码说明，'
        '为后续的编码实现提供依据。伪代码（Pseudocode）是一种算法描述语言，使用伪代码的目的是'
        '为了使被描述的算法可以容易地以任何一种编程语言实现。因此，伪代码必须结构清晰、代码简单、'
        '可读性好，并且类似自然语言。'
    )

    add_simple_table(
        document,
        ['层次', '技术栈', '职责', '主要文件'],
        [
            ('前端视图层', 'Vue 3 + Element Plus + Vue Router + Axios',
             '页面展示、用户交互、表单验证、路由导航、权限控制',
             'src/pages/*.vue, src/router/index.js, src/api/*.js'),
            ('后端控制层', 'Spring MVC Controller',
             '接收 HTTP 请求、参数绑定与校验、调用业务层、返回统一响应',
             'controller/*Controller.java'),
            ('后端业务层', 'Spring Service + @Transactional',
             '核心业务逻辑、数据校验、事务管理、统计计算',
             'service/*Service.java'),
            ('后端数据访问层', 'MyBatis-Plus Mapper + JdbcTemplate',
             '数据库 CRUD 操作、复杂 SQL 查询、聚合统计',
             'mapper/*Mapper.java'),
            ('数据实体层', 'JPA Entity + MyBatis-Plus 注解',
             '数据库表映射、字段定义、关联关系',
             'entity/*.java'),
        ],
        '表 3-1 系统五层架构划分',
    )

    add_body_paragraph(
        document,
        '以下各章节将按照功能模块逐一展开，每个模块按照视图层（前端页面）、控制层（Controller 类）、'
        '模型层（Service 类、Mapper 接口、Entity 类）的顺序给出详细设计说明和核心组件的伪代码。'
    )


# ── Section 4: Login & Auth ─────────────────────────────────────────

def add_login_module(document: DocumentObject) -> None:
    document.add_heading('4. 登录与权限管理模块的详细设计', level=1)

    add_body_paragraph(
        document,
        '登录与权限管理模块实现用户身份认证、Token 签发、角色识别与前端菜单权限控制。'
        '用户通过登录页提交用户名和密码，后端验证后返回 Token 及角色信息，前端将认证信息'
        '存储在 localStorage 中并据此控制路由访问和菜单显示。该模块的三层划分如表 4-1 所示。'
    )

    add_simple_table(
        document,
        ['层次', '组件/类', '说明'],
        [
            ('视图层', 'LoginPage.vue', '登录表单页面，采集用户名和密码'),
            ('视图层', 'MainLayout.vue', '主布局组件，根据角色动态显示菜单'),
            ('视图层', 'router/index.js', '路由配置与导航守卫，实现角色级访问控制'),
            ('控制层', 'AuthController', '处理 /auth/login 和 /auth/logout 请求'),
            ('模型层', 'AuthService', '用户认证、密码校验、角色查询逻辑'),
            ('模型层', 'SysUser', '用户实体，映射 sys_user 表'),
            ('模型层', 'SysRole', '角色实体，映射 sys_role 表'),
            ('模型层', 'SysUserRole', '用户角色关联实体，映射 sys_user_role 表'),
            ('模型层', 'SysUserMapper', '用户表数据访问'),
            ('模型层', 'SysRoleMapper', '角色表数据访问'),
            ('模型层', 'SysUserRoleMapper', '用户角色关联表数据访问'),
        ],
        '表 4-1 登录与权限管理模块的三层组件',
    )

    document.add_heading('4.1 视图层', level=2)
    add_body_paragraph(document, '登录模块视图层包含 1 个登录页面和 1 个主布局组件。')

    document.add_heading('4.1.1 LoginPage.vue 登录页面', level=3)
    add_body_paragraph(
        document,
        '登录页面是系统的入口页面，提供用户名和密码输入框及登录按钮。页面加载时预填演示账号信息，'
        '登录成功后调用 loginApi 获取 token 及用户信息，存储到 localStorage 并跳转到首页仪表盘。'
        '登录失败时通过 Element Plus 消息提示组件显示错误信息。登录页面的核心伪代码描述如下：'
    )

    add_code_block(document, [
        '/**',
        ' * @System: School Department Financial Management System',
        ' * @Component: LoginPage',
        ' * @Summary: 系统登录页面，采集用户名密码并完成认证',
        ' * @Framework: Vue 3 Composition API + Element Plus',
        ' */',
        '',
        '<template>',
        '  <div class="login-container">',
        '    <el-card class="login-card">',
        '      <h2>学校部门财务管理系统</h2>',
        '      <el-form ref="formRef" :model="loginForm" :rules="rules">',
        '        <el-form-item prop="username">',
        '          <el-input v-model="loginForm.username" placeholder="用户名" />',
        '        </el-form-item>',
        '        <el-form-item prop="password">',
        '          <el-input v-model="loginForm.password"',
        '            type="password" placeholder="密码" />',
        '        </el-form-item>',
        '        <el-form-item>',
        '          <el-button type="primary" @click="handleLogin">登录</el-button>',
        '        </el-form-item>',
        '      </el-form>',
        '    </el-card>',
        '  </div>',
        '</template>',
        '',
        '<script setup>',
        "import { reactive } from 'vue'",
        "import { useRouter } from 'vue-router'",
        "import { ElMessage } from 'element-plus'",
        "import { loginApi } from '@/api/finance'",
        "import { setAuth } from '@/utils/app'",
        '',
        'const router = useRouter()',
        'const loginForm = reactive({',
        "  username: 'admin',       // 预填演示账号",
        "  password: '123456',      // 预填演示密码",
        '})',
        '',
        'async function handleLogin() {',
        '  try {',
        '    const res = await loginApi(loginForm.username, loginForm.password)',
        '    setAuth({ token: res.token, username: res.username,',
        '              realName: res.realName, roles: res.roles })',
        "    ElMessage.success('登录成功')",
        "    router.push('/dashboard')",
        '  } catch (error) {',
        "    ElMessage.error(error.message || '登录失败')",
        '  }',
        '}',
        '</script>',
    ], '伪代码 4-1 LoginPage.vue 登录页面伪代码')

    document.add_heading('4.1.2 路由导航守卫', level=3)
    add_body_paragraph(
        document,
        '前端路由配置通过 Vue Router 的 beforeEach 导航守卫实现认证与授权控制。守卫逻辑分为三层：'
        '未认证用户重定向到登录页；已认证用户访问登录页重定向到仪表盘；根据用户角色检查路由所需权限。'
    )

    add_code_block(document, [
        '/**',
        ' * @Router: 路由导航守卫',
        ' * @Summary: 在每次路由切换前检查认证状态和角色权限',
        ' */',
        '',
        '// 路由配置示例（角色元信息定义）',
        'const routes = [',
        "  { path: '/login', component: LoginPage, meta: { public: true } },",
        "  { path: '/dashboard', component: DashboardPage,",
        "    meta: { roles: ['ADMIN', 'FINANCE'] } },",
        "  { path: '/users', component: UserPage, meta: { roles: ['ADMIN'] } },",
        '  // ... 其他路由',
        ']',
        '',
        '// 全局前置守卫',
        'router.beforeEach((to, from, next) => {',
        '  const auth = getAuth()  // 从 localStorage 读取认证信息',
        '  const isAuth = auth && auth.token',
        '',
        '  // 规则1: 公开页面已认证则重定向到仪表盘',
        '  if (to.meta.public) {',
        "    return isAuth ? next('/dashboard') : next()",
        '  }',
        '',
        '  // 规则2: 非公开页面需要认证',
        "  if (!isAuth) return next('/login')",
        '',
        '  // 规则3: 角色授权检查',
        '  const requiredRoles = to.meta.roles',
        '  if (requiredRoles && requiredRoles.length > 0) {',
        '    const hasRole = requiredRoles.some(r => (auth.roles||[]).includes(r))',
        "    if (!hasRole) return next('/dashboard')",
        '  }',
        '',
        '  next()',
        '})',
    ], '伪代码 4-2 路由导航守卫伪代码')

    document.add_heading('4.1.3 主布局与菜单控制', level=3)
    add_body_paragraph(
        document,
        'MainLayout.vue 是系统的主布局组件，包含顶部导航栏、左侧菜单栏和中间内容区。菜单项根据'
        '当前用户角色动态过滤显示：管理员可访问全部菜单，财务员仅显示仪表盘、收支记录和报表。'
    )

    document.add_heading('4.2 控制层', level=2)
    add_body_paragraph(
        document,
        '登录模块控制层由 AuthController 类实现，提供两个 REST 接口：/auth/login（登录）'
        '和 /auth/logout（退出）。'
    )

    add_code_block(document, [
        '/**',
        ' * @Class:   AuthController',
        ' * @Summary: 认证控制器，处理登录和退出请求',
        ' * @BasePath: /auth',
        ' */',
        '@RestController',
        '@RequestMapping("/auth")',
        'public class AuthController {',
        '    private final AuthService authService;',
        '',
        '    /**',
        '     * POST /auth/login',
        '     * 输入: LoginRequest { username, password }',
        '     * 输出: LoginResponse { token, username, realName, roles }',
        '     */',
        '    @PostMapping("/login")',
        '    public ApiResponse<LoginResponse> login(',
        '            @Valid @RequestBody LoginRequest request) {',
        '        return ApiResponse.success(authService.login(request))',
        '    }',
        '',
        '    /** POST /auth/logout — 用户退出 */',
        '    @PostMapping("/logout")',
        '    public ApiResponse<Map<String, String>> logout() {',
        '        return ApiResponse.success(Map.of("message", "退出成功"))',
        '    }',
        '}',
    ], '伪代码 4-3 AuthController.java 伪代码')

    document.add_heading('4.3 模型层', level=2)

    document.add_heading('4.3.1 AuthService 认证服务', level=3)
    add_body_paragraph(
        document,
        'AuthService 是认证模块的核心业务类，负责用户名校验、密码验证（支持明文/BCrypt/演示密码'
        '三种模式）、用户状态检查、角色信息加载和 Token 生成。'
    )

    add_code_block(document, [
        '/**',
        ' * @Class:   AuthService',
        ' * @Summary: 认证服务，处理登录：用户校验、密码验证、角色加载、Token 生成',
        ' */',
        '@Service',
        'public class AuthService {',
        '    private final SysUserMapper userMapper;',
        '    private final SysUserRoleMapper userRoleMapper;',
        '    private final SysRoleMapper roleMapper;',
        '    private final PasswordEncoder passwordEncoder;',
        '',
        '    public LoginResponse login(LoginRequest request) {',
        '        // Step 1: 查询用户',
        '        SysUser user = userMapper.selectOne(',
        '            new LambdaQueryWrapper<SysUser>()',
        '                .eq(SysUser::getUsername, request.getUsername()))',
        '        if (user == null) throw new BusinessException("用户名或密码错误")',
        '',
        '        // Step 2: 检查账号状态',
        '        if (user.getStatus() == 0)',
        '            throw new BusinessException("账号已被停用，请联系管理员")',
        '',
        '        // Step 3: 密码校验（兼容明文/演示/BCrypt 三种模式）',
        '        if (!passwordMatches(request.getPassword(), user.getPassword()))',
        '            throw new BusinessException("用户名或密码错误")',
        '',
        '        // Step 4: 加载角色信息',
        '        List<String> roles = new ArrayList<>()',
        '        List<SysUserRole> userRoles = userRoleMapper.selectList(',
        '            new LambdaQueryWrapper<SysUserRole>()',
        '                .eq(SysUserRole::getUserId, user.getId()))',
        '        for (SysUserRole ur : userRoles) {',
        '            SysRole role = roleMapper.selectById(ur.getRoleId())',
        '            if (role != null) roles.add(role.getRoleCode())',
        '        }',
        '',
        '        // Step 5: 生成 Token 并返回',
        '        String token = UUID.randomUUID().toString()',
        '        return new LoginResponse(token, user.getUsername(),',
        '            user.getRealName(), roles)',
        '    }',
        '}',
    ], '伪代码 4-4 AuthService.java 伪代码')

    document.add_heading('4.3.2 DTO/VO 类设计', level=3)
    add_code_block(document, [
        'public record LoginRequest(',
        '    @NotBlank(message = "用户名不能为空") String username,',
        '    @NotBlank(message = "密码不能为空")   String password',
        ') {}',
        '',
        'public record LoginResponse(',
        '    String token,       // UUID 令牌',
        '    String username,    // 登录用户名',
        '    String realName,    // 真实姓名',
        '    List<String> roles  // 角色编码列表',
        ') {}',
    ], '伪代码 4-5 登录模块 DTO/VO 伪代码')


# ── Section 5: User Management ──────────────────────────────────────

def add_user_module(document: DocumentObject) -> None:
    document.add_heading('5. 用户管理模块的详细设计', level=1)

    add_body_paragraph(
        document,
        '用户管理模块面向管理员提供系统用户的增删改查和状态管理功能。管理员可通过用户管理页面'
        '创建新用户、修改已有用户信息、以及停用/启用用户账号。'
    )

    add_simple_table(
        document,
        ['层次', '组件/类', '说明'],
        [
            ('视图层', 'UserPage.vue', '用户管理页面，含列表、搜索、弹窗表单'),
            ('视图层', 'PageCard.vue', '通用卡片容器组件'),
            ('控制层', 'UserController', '处理 /users 的 CRUD 请求'),
            ('模型层', 'UserService', '用户创建、更新、状态变更、唯一性校验'),
            ('模型层', 'UserSaveRequest', '用户保存请求 DTO，含 @Valid 校验'),
            ('模型层', 'UserVO', '用户视图对象，含角色名称'),
            ('模型层', 'SysUserMapper', '用户表数据访问'),
            ('模型层', 'SysUserRoleMapper', '用户角色关联表数据访问'),
            ('模型层', 'SysRoleMapper', '角色表数据访问'),
        ],
        '表 5-1 用户管理模块的三层组件',
    )

    document.add_heading('5.1 视图层', level=2)
    add_body_paragraph(
        document,
        'UserPage.vue 采用列表+弹窗的交互模式。页面顶部提供搜索输入框和新增按钮，'
        '主体区域展示用户列表（用户名、真实姓名、电话、角色、状态），每行提供编辑和状态切换按钮。'
        '编辑/新增操作通过 Element Plus Dialog 弹窗完成。'
    )

    document.add_heading('5.2 控制层', level=2)
    add_code_block(document, [
        '@RestController',
        '@RequestMapping("/users")',
        'public class UserController {',
        '    private final UserService userService;',
        '',
        '    @GetMapping              // GET  /users — 查询全部用户',
        '    @PostMapping             // POST /users — 新增用户',
        '    @PutMapping("/{id}")      // PUT  /users/{id} — 修改用户',
        '    @PatchMapping("/{id}/status") // PATCH — 修改状态',
        '}',
    ], '伪代码 5-1 UserController.java 伪代码')

    document.add_heading('5.3 模型层', level=2)
    add_code_block(document, [
        '@Service',
        'public class UserService {',
        '    private final SysUserMapper userMapper;',
        '    private final SysUserRoleMapper userRoleMapper;',
        '    private final SysRoleMapper roleMapper;',
        '    private final PasswordEncoder passwordEncoder;',
        '',
        '    public List<UserVO> list() {',
        '        // 查询所有用户，通过关联表填充角色名称',
        '        // 流式处理：user -> userRole -> roleName -> UserVO',
        '        ...',
        '    }',
        '',
        '    @Transactional',
        '    public UserVO create(UserSaveRequest req) {',
        '        // 1) 密码非空校验  2) 用户名唯一性检查',
        '        // 3) BCrypt 加密密码, 插入用户',
        '        // 4) 根据 roleCode 查角色, 插入 sys_user_role',
        '    }',
        '',
        '    @Transactional',
        '    public UserVO update(Long id, UserSaveRequest req) {',
        '        // 1) 获取用户，不存在抛异常',
        '        // 2) 用户名唯一性检查（排除自身）',
        '        // 3) 更新字段，有密码则加密  4) 更新角色关联',
        '    }',
        '',
        '    public void updateStatus(Long id, Integer status) {',
        '        SysUser u = getUser(id); u.setStatus(status); userMapper.updateById(u)',
        '    }',
        '}',
    ], '伪代码 5-2 UserService.java 伪代码')


# ── Section 6: Department Management ─────────────────────────────────

def add_department_module(document: DocumentObject) -> None:
    document.add_heading('6. 部门管理模块的详细设计', level=1)
    add_body_paragraph(
        document,
        '部门管理模块负责维护 dept_info 表中的部门基础信息（编号、名称、负责人、电话）。'
        '采用停用优先策略，被收支记录引用的部门不执行物理删除。'
    )
    add_simple_table(
        document,
        ['层次', '组件/类', '说明'],
        [
            ('视图层', 'DepartmentPage.vue', '部门管理页面，列表+弹窗模式'),
            ('控制层', 'DepartmentController', '处理 /departments CRUD 请求'),
            ('模型层', 'DepartmentService', '增改查、状态管理、编号唯一性校验'),
            ('模型层', 'DeptInfo / DeptInfoMapper', '实体与数据访问'),
        ],
        '表 6-1 部门管理模块的三层组件',
    )
    document.add_heading('6.1 控制层', level=2)
    add_code_block(document, [
        '@RestController @RequestMapping("/departments")',
        'public class DepartmentController {',
        '    // GET /departments, POST /departments,',
        '    // PUT /departments/{id}, PATCH /departments/{id}/status',
        '}',
    ], '伪代码 6-1 DepartmentController.java')
    document.add_heading('6.2 模型层', level=2)
    add_code_block(document, [
        '@Service',
        'public class DepartmentService {',
        '    // list() - 按 ID 升序查询全部部门',
        '    // create(req) - 编号唯一性检查 + 插入',
        '    // update(id, req) - 编号唯一性(排除自身) + 更新',
        '    // updateStatus(id, status) - 状态切换',
        '    // getEnabled(id) - 获取有效部门，停用则抛 BusinessException',
        '}',
    ], '伪代码 6-2 DepartmentService.java')


# ── Section 7: Account Management ───────────────────────────────────

def add_account_module(document: DocumentObject) -> None:
    document.add_heading('7. 账户管理模块的详细设计', level=1)
    add_body_paragraph(
        document,
        '账户管理模块维护 account_info 表（账户编号、名称、开户单位、说明）。'
        '系统在录入收支记录时仅允许选择有效账户，保证业务数据一致性。'
    )
    add_simple_table(
        document,
        ['层次', '组件/类', '说明'],
        [
            ('视图层', 'AccountPage.vue', '账户管理页面，交互模式与部门管理一致'),
            ('控制层', 'AccountController', '处理 /accounts CRUD 请求'),
            ('模型层', 'AccountService', '增改查、状态管理、编号唯一性校验'),
            ('模型层', 'AccountInfo / AccountInfoMapper', '实体与数据访问'),
        ],
        '表 7-1 账户管理模块的三层组件',
    )
    add_code_block(document, [
        '@RestController @RequestMapping("/accounts")',
        'public class AccountController { /* CRUD 结构同部门管理 */ }',
        '',
        '@Service',
        'public class AccountService {',
        '    // list(), create(req), update(id, req),',
        '    // updateStatus(id, status), getEnabled(id)',
        '}',
    ], '伪代码 7-1 账户管理控制层与模型层伪代码')


# ── Section 8: Category Management ──────────────────────────────────

def add_category_modules(document: DocumentObject) -> None:
    document.add_heading('8. 类别管理模块的详细设计', level=1)
    add_body_paragraph(
        document,
        '类别管理分为收入类别（income_category）和支出类别（expense_category）两个对称子模块，'
        '分别维护各自类别的编号、名称和状态。类别数据作为报表统计的重要维度。'
    )
    document.add_heading('8.1 控制层', level=2)
    add_code_block(document, [
        '@RestController @RequestMapping("/income-categories")',
        'public class IncomeCategoryController { /* CRUD */ }',
        '',
        '@RestController @RequestMapping("/expense-categories")',
        'public class ExpenseCategoryController { /* CRUD 对称结构 */ }',
    ], '伪代码 8-1 类别控制器')
    document.add_heading('8.2 模型层', level=2)
    add_code_block(document, [
        '@Service',
        'public class CategoryService {',
        '    // 收入类别: listIncomeCategories(), createIncomeCategory(),',
        '    //   updateIncomeCategory(), updateIncomeCategoryStatus(),',
        '    //   getEnabledIncomeCategory()',
        '    // 支出类别: 方法结构完全对称',
        '}',
    ], '伪代码 8-2 CategoryService.java')


# ── Section 9: Income Record ────────────────────────────────────────

def add_income_record_module(document: DocumentObject) -> None:
    document.add_heading('9. 收入记录管理模块的详细设计', level=1)
    add_body_paragraph(
        document,
        '收入记录管理是核心业务模块，负责收入记录的录入、修改、查询和停用。每笔记录关联部门、'
        '账户和收入类别，系统自动生成 IN{yyyyMMdd}{NNN} 格式的唯一编号。'
    )

    add_simple_table(
        document,
        ['层次', '组件/类', '说明'],
        [
            ('视图层', 'IncomePage.vue', '收入记录页，含查询筛选、列表和录入弹窗'),
            ('控制层', 'IncomeController', '处理 /incomes CRUD 请求'),
            ('模型层', 'RecordService', '增改查、外键校验、编号生成'),
            ('模型层', 'IncomeRecord / Mapper', '实体与数据访问'),
            ('模型层', 'FinanceRecordSaveRequest / VO', 'DTO 与视图对象'),
        ],
        '表 9-1 收入记录管理模块的三层组件',
    )

    document.add_heading('9.1 视图层', level=2)
    add_body_paragraph(
        document,
        'IncomePage.vue 上部提供部门/账户/类别下拉筛选和日期区间选择，中部展示收入记录列表，'
        '点击新增/编辑弹出 Dialog 表单。支出记录页面 ExpensePage.vue 结构完全对称。'
    )

    document.add_heading('9.2 控制层', level=2)
    add_code_block(document, [
        '@RestController @RequestMapping("/incomes")',
        'public class IncomeController {',
        '    // GET /incomes?deptId=&accountId=&categoryId=&startDate=&endDate=',
        '    //    → 多条件筛选，JdbcTemplate JOIN 查询',
        '    // POST /incomes → 创建收入记录，自动生成 IN 编号',
        '    // PUT /incomes/{id} → 修改收入记录',
        '    // PATCH /incomes/{id}/status → 修改状态',
        '}',
    ], '伪代码 9-1 IncomeController.java')

    document.add_heading('9.3 模型层', level=2)
    add_code_block(document, [
        '@Service',
        'public class RecordService {',
        '    // ── 收入记录 ──',
        '    public Map<String,Object> listIncomes(filters...) {',
        '        // JdbcTemplate JOIN income_record + dept_info',
        '        //   + account_info + income_category',
        '        // 动态 WHERE 条件拼接',
        '    }',
        '',
        '    @Transactional',
        '    public FinanceRecordVO createIncome(FinanceRecordSaveRequest req) {',
        '        // Step 1: validateIncomeRequest(req)',
        '        //   → departmentService.getEnabled(req.deptId)',
        '        //   → accountService.getEnabled(req.accountId)',
        '        //   → categoryService.getEnabledIncomeCategory(req.categoryId)',
        '        // Step 2: generateRecordNo("IN")',
        '        //   → IN{yyyyMMdd}{3位序号}, 从 001 开始递增',
        '        // Step 3: 构建 IncomeRecord 并 insert',
        '        // Step 4: 返回 FinanceRecordVO（含关联表名称）',
        '    }',
        '',
        '    // updateIncome(), updateIncomeStatus() 结构类似',
        '',
        '    // ── 支出记录 ──',
        '    // listExpenses(), createExpense(), updateExpense() 结构对称',
        '    // 编号前缀为 "EX"，关联 expense_category 表',
        '}',
    ], '伪代码 9-2 RecordService.java 核心逻辑')


# ── Section 10: Expense Record ──────────────────────────────────────

def add_expense_record_module(document: DocumentObject) -> None:
    document.add_heading('10. 支出记录管理模块的详细设计', level=1)
    add_body_paragraph(
        document,
        '支出记录管理模块与收入记录完全对称。编号格式为 EX{yyyyMMdd}{NNN}，关联支出类别表。'
        '高度对称的设计有利于前端页面结构统一、后端校验逻辑复用和报表计算口径稳定。'
    )
    add_simple_table(
        document,
        ['层次', '组件/类', '说明'],
        [
            ('视图层', 'ExpensePage.vue', '与 IncomePage.vue 结构对称'),
            ('控制层', 'ExpenseController', '/expenses CRUD，与 IncomeController 对称'),
            ('模型层', 'RecordService(复用)', '调用 createExpense/updateExpense 等方法族'),
            ('模型层', 'ExpenseRecord / Mapper', '实体与数据访问'),
        ],
        '表 10-1 支出记录管理模块的三层组件',
    )
    add_code_block(document, [
        '@RestController @RequestMapping("/expenses")',
        'public class ExpenseController {',
        '    // GET /expenses?deptId=&... → 多条件筛选',
        '    // POST /expenses → 创建, PATCH /expenses/{id}/status → 状态',
        '}',
    ], '伪代码 10-1 ExpenseController.java')


# ── Section 11: Report ──────────────────────────────────────────────

def add_report_module(document: DocumentObject) -> None:
    document.add_heading('11. 财务统计与报表模块的详细设计', level=1)
    add_body_paragraph(
        document,
        '报表模块通过 /reports 接口组提供五种统计能力：周报、月报、年报、自定义区间报表和首页仪表盘。'
        '统计逻辑只读取 status=1 的有效记录，按时间区间汇总收入、支出、结余，并分维度统计。'
    )

    add_simple_table(
        document,
        ['层次', '组件/类', '说明'],
        [
            ('视图层', 'ReportPage.vue / DashboardPage.vue', '报表中心与首页仪表盘'),
            ('控制层', 'ReportController', '/reports/dashboard, weekly, monthly, yearly, custom'),
            ('模型层', 'ReportService', '时间区间解析、聚合统计、分维度汇总与趋势计算'),
        ],
        '表 11-1 报表模块的三层组件',
    )

    document.add_heading('11.1 控制层', level=2)
    add_code_block(document, [
        '@RestController @RequestMapping("/reports")',
        'public class ReportController {',
        '    // GET /reports/dashboard — 当月汇总',
        '    // GET /reports/weekly?week=2026-W23&deptId=1',
        '    // GET /reports/monthly?month=2026-06&deptId=1',
        '    // GET /reports/yearly?year=2026&deptId=1',
        '    // GET /reports/custom?startDate=&endDate=&deptId=',
        '}',
    ], '伪代码 11-1 ReportController.java')

    document.add_heading('11.2 模型层', level=2)
    add_code_block(document, [
        '@Service',
        'public class ReportService {',
        '    // 核心方法 buildSummary(startDate, endDate, deptId):',
        '    //   Step 1: sumAmount("income_record", "amount", ...)',
        '    //           → SELECT COALESCE(SUM(amount),0) WHERE status=1',
        '    //   Step 2: sumAmount("expense_record", "amount", ...)',
        '    //   Step 3: balance = totalIncome - totalExpense',
        '    //   Step 4: buildCategoryStats — UNION ALL 收入+支出类别',
        '    //   Step 5: buildDepartmentStats — LEFT JOIN 子查询',
        '    //   Step 6: buildTrendStats — 按日/月聚合，填满所有日期',
        '    //           >62天或整年 → 按月(DATE_FORMAT)',
        '    //           否则 → 按日',
        '    // 返回 ReportSummaryVO',
        '}',
    ], '伪代码 11-2 ReportService.java 核心统计逻辑')


# ── Section 12: Backup ──────────────────────────────────────────────

def add_backup_module(document: DocumentObject) -> None:
    document.add_heading('12. 数据备份与恢复模块的详细设计', level=1)
    add_body_paragraph(
        document,
        '备份模块管理数据库备份文件的完整生命周期：创建备份时生成包含全部 10 张表数据的 .sql 文件；'
        '下载时验证文件存在后以流返回；恢复时读取 SQL 逐句执行并记录恢复时间。'
    )

    add_simple_table(
        document,
        ['层次', '组件/类', '说明'],
        [
            ('视图层', 'BackupPage.vue', '备份列表、创建按钮、下载/恢复操作（恢复需二次确认）'),
            ('控制层', 'BackupController', '/backups 列表/创建/下载/恢复'),
            ('模型层', 'BackupService', 'SQL 构建、文件 IO、恢复执行'),
            ('模型层', 'BackupRecord / Mapper', '备份元数据实体与数据访问'),
        ],
        '表 12-1 备份模块的三层组件',
    )

    document.add_heading('12.1 控制层', level=2)
    add_code_block(document, [
        '@RestController @RequestMapping("/backups")',
        'public class BackupController {',
        '    // GET /backups — 列表',
        '    // POST /backups — 创建备份',
        '    // GET /backups/{id}/download — 下载（ResponseEntity<Resource>）',
        '    // POST /backups/{id}/restore — 恢复（更新 restoredAt）',
        '}',
    ], '伪代码 12-1 BackupController.java')

    document.add_heading('12.2 模型层', level=2)
    add_code_block(document, [
        '@Service',
        'public class BackupService {',
        '    @Value("${app.backup-dir:./backups}")',
        '    private String backupDir;',
        '',
        '    @PostConstruct',
        '    public void init() { Files.createDirectories(Path.of(backupDir)) }',
        '',
        '    @Transactional',
        '    public BackupFileVO createBackup(BackupCreateRequest req) {',
        '        // 1) 生成文件名: finance_backup_{yyyyMMddHHmmss}.sql',
        '        // 2) buildBackupContent():',
        '        //    a) SET FOREIGN_KEY_CHECKS = 0',
        '        //    b) DELETE 语序(按反向依赖):',
        '        //       expense_record → income_record → backup_record',
        '        //       → sys_user_role → expense_category → income_category',
        '        //       → account_info → dept_info → sys_user → sys_role',
        '        //    c) INSERT 语序(按正向依赖): 遍历 10 张表全量数据',
        '        //    d) SET FOREIGN_KEY_CHECKS = 1',
        '        // 3) 写入磁盘 + 保存 backup_record 元数据',
        '    }',
        '',
        '    public Resource download(Long id) {',
        '        // 验证记录和文件存在 → 返回 FileSystemResource',
        '    }',
        '',
        '    @Transactional',
        '    public BackupFileVO restore(Long id) {',
        '        // 读取 SQL → 按;分割 → 逐句 jdbcTemplate.execute()',
        '        // → 更新 backup_record.restoredAt',
        '    }',
        '}',
    ], '伪代码 12-2 BackupService.java')


# ── Section 13: Frontend Client ─────────────────────────────────────

def add_client_module(document: DocumentObject) -> None:
    document.add_heading('13. 客户端模块的详细设计', level=1)
    add_body_paragraph(
        document,
        '客户端模块是整个前端 SPA 应用的框架层，包含项目初始化、路由配置、HTTP 通信封装、'
        '认证工具函数、全局样式和通用组件。基于 Vue 3 + Vite + Element Plus 构建，通过 Axios '
        '与后端 REST API 通信。'
    )

    add_simple_table(
        document,
        ['组件/文件', '类型', '说明'],
        [
            ('main.js', '入口', '创建 Vue 应用，注册 Element Plus、Router、样式'),
            ('App.vue', '根组件', '仅含 <router-view />，作为路由出口'),
            ('router/index.js', '路由', '定义全部路由、元信息和导航守卫'),
            ('api/http.js', 'HTTP', 'Axios 实例，请求/响应拦截器'),
            ('api/finance.js', 'API', '封装全部后端 API 调用函数'),
            ('utils/app.js', '工具', '认证信息存取（localStorage）'),
            ('layout/MainLayout.vue', '布局', '顶部导航+侧边菜单+内容区'),
            ('components/PageCard.vue', '组件', '通用卡片容器'),
            ('styles/global.css', '样式', '全局样式、渐变色背景'),
        ],
        '表 13-1 客户端模块文件清单',
    )

    document.add_heading('13.1 视图层 — HTTP 通信模块（http.js）', level=2)
    add_code_block(document, [
        '// Axios 实例创建',
        "const http = axios.create({",
        "  baseURL: 'http://localhost:8080',",
        '  timeout: 15000,',
        '})',
        '',
        '// 请求拦截器：自动附加 Authorization: Bearer {token}',
        'http.interceptors.request.use(config => {',
        '  const auth = getAuth()',
        '  if (auth && auth.token)',
        '    config.headers.Authorization = `Bearer ${auth.token}`',
        '  return config',
        '})',
        '',
        '// 响应拦截器：统一解包 ApiResponse 信封',
        'http.interceptors.response.use(',
        '  response => {',
        '    const body = response.data  // { code, message, data }',
        '    if (body.code === 200) return body.data  // 成功 → 返回 data',
        "    ElMessage.error(body.message || '请求失败')",
        '    return Promise.reject(new Error(body.message))',
        '  },',
        '  error => {',
        "    ElMessage.error(error.message || '网络错误')",
        '    return Promise.reject(error)',
        '  }',
        ')',
    ], '伪代码 13-1 http.js Axios 封装')

    document.add_heading('13.2 视图层 — API 调用模块（finance.js）', level=2)
    add_body_paragraph(
        document,
        'finance.js 封装全部后端 API 调用：认证接口（login/logout）、通用 CRUD 辅助函数'
        '（fetchList/createItem/updateItem/updateItemStatus）、各模块 API 函数以及常量。'
        '前端页面通过调用这些函数与后端进行数据交互，无需直接使用 Axios。'
    )

    document.add_heading('13.3 视图层 — 认证工具（app.js）', level=2)
    add_code_block(document, [
        "const AUTH_KEY = 'finance-auth'",
        '',
        '// 存储: localStorage.setItem(AUTH_KEY, JSON.stringify(auth))',
        '// auth = { token, username, realName, roles }',
        '',
        'export function setAuth(auth) { localStorage.setItem(...) }',
        'export function getAuth() { return JSON.parse(localStorage.getItem(...)) }',
        'export function clearAuth() { localStorage.removeItem(...) }',
    ], '伪代码 13-2 app.js 认证工具')


# ── Section 14: Reused Modules ──────────────────────────────────────

def add_reused_modules(document: DocumentObject) -> None:
    document.add_heading('14. 复用的模块', level=1)
    add_body_paragraph(
        document,
        '在整个开发过程中，系统采用了复用原则，将通用功能抽取为独立模块供各处引用。'
    )

    document.add_heading('14.1 统一 API 响应格式', level=2)
    add_body_paragraph(
        document,
        '后端所有接口统一返回 ApiResponse<T> 信封对象（code + message + data），'
        '前端 Axios 响应拦截器统一解包，各业务接口只需关注 data 内容。'
    )
    add_code_block(document, [
        'public record ApiResponse<T>(int code, String message, T data) {',
        '    public static <T> ApiResponse<T> success(T data) {',
        '        return new ApiResponse<>(200, "success", data)',
        '    }',
        '    public static <T> ApiResponse<T> failure(String message) {',
        '        return new ApiResponse<>(500, message, null)',
        '    }',
        '}',
    ], '伪代码 14-1 ApiResponse.java')

    document.add_heading('14.2 统一异常处理', level=2)
    add_body_paragraph(
        document,
        'GlobalExceptionHandler（@RestControllerAdvice）将 BusinessException 和校验异常'
        '统一转换为 ApiResponse 格式返回，避免异常信息泄露到前端。'
    )

    document.add_heading('14.3 MyBatis-Plus 基础 CRUD', level=2)
    add_body_paragraph(
        document,
        '所有 Mapper 接口继承 BaseMapper<T>，自动获得 insert/updateById/selectById/selectList 等'
        '20+ 内置方法。简单 CRUD 操作无需编写 SQL，复杂查询使用 JdbcTemplate 编写原生 SQL。'
    )

    document.add_heading('14.4 其他复用', level=2)
    add_bullet_paragraph(document, 'PageCard.vue — 前端通用卡片容器组件，各页面统一使用')
    add_bullet_paragraph(document, 'WebConfig — 全局 CORS 配置，允许前端跨域访问')
    add_bullet_paragraph(document, 'SecurityConfig — BCryptPasswordEncoder Bean，密码安全存储')


# ── Section 15: Configuration ───────────────────────────────────────

def add_configuration(document: DocumentObject) -> None:
    document.add_heading('15. 配置文件', level=1)

    document.add_heading('15.1 application.yml 后端配置', level=2)
    add_code_block(document, [
        '# application.yml',
        'server:',
        '  port: 8080',
        'spring:',
        '  datasource:',
        '    url: jdbc:mysql://localhost:3306/finance_management_system',
        '      ?useUnicode=true&characterEncoding=utf8&serverTimezone=Asia/Shanghai',
        '    username: root',
        '    password: 123456',
        'mybatis-plus:',
        '  configuration:',
        '    map-underscore-to-camel-case: true',
        'app:',
        '  backup-dir: ./backups',
    ], '伪代码 15-1 application.yml')

    document.add_heading('15.2 前端构建配置', level=2)
    add_body_paragraph(
        document,
        'vite.config.js 配置开发服务器端口 5173 和 /api 代理到 localhost:8080。'
        'package.json 声明核心依赖：vue ^3.4.27, vue-router ^4.4.0, element-plus ^2.7.5, '
        'axios ^1.7.2, vite ^5.2.0。'
    )

    document.add_heading('15.3 后端依赖与数据库脚本', level=2)
    add_body_paragraph(
        document,
        'pom.xml 声明核心依赖：spring-boot-starter-web 3.3.0, mybatis-plus-spring-boot3-starter 3.5.7, '
        'mysql-connector-j, spring-security-crypto, lombok。'
        '财务管理系统.sql（11-数据库/）创建 10 张表并插入种子数据，含 2 个演示用户和示例业务记录。'
    )


# ── Section 16: Conclusion ──────────────────────────────────────────

def add_conclusion(document: DocumentObject) -> None:
    document.add_heading('16. 结论', level=1)
    add_body_paragraph(
        document,
        '本文档按照前后端分离的架构，对学校部门财务管理系统的 10 个功能模块和客户端框架进行了'
        '详细的层次化设计描述。每个模块均从视图层（Vue 页面与组件）、控制层（Spring MVC Controller）'
        '和模型层（Service + Mapper + Entity）三个维度展开，并提供了核心组件的伪代码，为后续的'
        '编码、测试和维护工作提供了清晰的技术依据。'
    )
    add_body_paragraph(
        document,
        '本详细设计方案充分利用了 Spring Boot 3 自动配置、MyBatis-Plus 内置 CRUD、Vue 3 组合式 API '
        '和 Element Plus 组件库等现代框架能力，在保证功能完整性的同时控制了代码复杂度和开发周期。'
        '系统模块划分清晰、层次职责明确，便于后续的功能扩展和团队协作。'
    )


# ── Document assembly ────────────────────────────────────────────────

def build_document() -> DocumentObject:
    document = Document()
    set_document_defaults(document)

    create_cover(document)
    add_page_break(document)
    add_version_record(document)
    add_page_break(document)

    document.add_heading('目 录', level=1)
    toc_paragraph = document.add_paragraph()
    toc_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_toc(toc_paragraph)

    intro_section = document.add_section(WD_SECTION_START.NEW_PAGE)
    restart_page_numbering(intro_section)
    set_footer_page_number(intro_section)

    add_introduction(document)
    add_system_overview(document)
    add_detailed_design_overview(document)
    add_login_module(document)
    add_user_module(document)
    add_department_module(document)
    add_account_module(document)
    add_category_modules(document)
    add_income_record_module(document)
    add_expense_record_module(document)
    add_report_module(document)
    add_backup_module(document)
    add_client_module(document)
    add_reused_modules(document)
    add_configuration(document)
    add_conclusion(document)

    return document


def main() -> None:
    document = build_document()
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(OUTPUT_PATH))
    print(f'Generated: {OUTPUT_PATH}')


if __name__ == '__main__':
    main()
